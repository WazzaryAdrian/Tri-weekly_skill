#!/usr/bin/env python3
"""按 Issue 7 真实规格生成三周报 .docx。
用法: python generate_docx.py input.json output.docx  (结构见 references/issue_schema.example.json)
规格(取自 Volume_2_Issue_7 实测):
  文件开头直接是“目录”(无标题行/日期); 页面 Letter; 边距 上下1" 左右1.25" 页眉脚0.5"; 宋体/Cambria。
  正文(摘要/【原文信息】) 12pt(sz24) 行距1.16(line278) 段后8pt(160); 摘要两端对齐, 【原文信息】与标题左对齐。
  板块标题 加粗14pt(黑) 行距1.5 段前后10pt; 条目标题 加粗12pt(黑) 行距1.16 段后8pt。
  每个子模块(条目标题→摘要→【原文信息】)之后另插一个空行段(12pt/1.16/段后8pt)。
  目录: 一级 加粗11pt(sz22)、二级 常规10pt(sz20), 字号直接写死在文字上(不依赖目录样式);
        点引线右对齐页码用 PAGEREF 域(打开 Word 后 F9/更新域 填正确页码)。
"""
import json, sys
from docx import Document
from docx.shared import Pt, Twips, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

EA="宋体"; LATIN="Cambria"; CONTENT_W=8640; BODY=1.16

def font_run(run, size_pt, bold=False):
    run.font.size=Pt(size_pt); run.font.bold=bold; run.font.name=LATIN
    rpr=run._element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rpr.insert(0,rf)
    rf.set(qn('w:ascii'),LATIN); rf.set(qn('w:hAnsi'),LATIN); rf.set(qn('w:eastAsia'),EA)

def style_fonts(style, size_pt, bold):
    f=style.font; f.name=LATIN; f.size=Pt(size_pt); f.bold=bold; f.color.rgb=RGBColor(0,0,0)
    rpr=style.element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rpr.insert(0,rf)
    rf.set(qn('w:ascii'),LATIN); rf.set(qn('w:hAnsi'),LATIN); rf.set(qn('w:eastAsia'),EA)

def get_or_add_style(doc, name, base="Normal"):
    try: return doc.styles[name]
    except KeyError:
        st=doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH); st.base_style=doc.styles[base]; return st

def dotted_tab(style):
    style.paragraph_format.tab_stops.add_tab_stop(Twips(CONTENT_W), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

def para(doc, mult=BODY, before=0, after=0, align=None, style=None):
    p=doc.add_paragraph(style=style); pf=p.paragraph_format
    pf.line_spacing_rule=WD_LINE_SPACING.MULTIPLE; pf.line_spacing=mult
    pf.space_before=Twips(before); pf.space_after=Twips(after)
    if align is not None: p.alignment=align
    return p

def _fld(t):
    e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),t); return e

def heading_bm(doc, style_name, text, size_pt, bm_name, bm_id):
    """标题段 + 包一个书签(供目录 PAGEREF 指向)。"""
    p=doc.add_paragraph(style=style_name)
    st=OxmlElement('w:bookmarkStart'); st.set(qn('w:id'),str(bm_id)); st.set(qn('w:name'),bm_name); p._p.append(st)
    font_run(p.add_run(text), size_pt, True)
    en=OxmlElement('w:bookmarkEnd'); en.set(qn('w:id'),str(bm_id)); p._p.append(en)
    return p

def toc_entry(doc, style_name, text, bm_name, size_pt, bold):
    """目录行: 文字(字号写死) + 制表(点引线) + PAGEREF 页码域。"""
    p=doc.add_paragraph(style=style_name)
    font_run(p.add_run(text), size_pt, bold)
    font_run(p.add_run('\t'), size_pt, bold)
    p.add_run()._element.append(_fld('begin'))
    it=OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text=f' PAGEREF {bm_name} \\h '
    p.add_run()._element.append(it)
    p.add_run()._element.append(_fld('separate'))
    font_run(p.add_run('1'), size_pt, bold)          # 占位页码, F9 后更新
    p.add_run()._element.append(_fld('end'))
    return p

def main():
    data=json.load(open(sys.argv[1],encoding='utf-8'))
    out=sys.argv[2] if len(sys.argv)>2 else "三周报_output.docx"
    doc=Document()
    s=doc.sections[0]
    s.page_width=Twips(12240); s.page_height=Twips(15840)
    s.top_margin=Twips(1440); s.bottom_margin=Twips(1440)
    s.left_margin=Twips(1800); s.right_margin=Twips(1800)
    s.header_distance=Twips(720); s.footer_distance=Twips(720)

    h1=doc.styles['Heading 1']; style_fonts(h1,14,True)
    h1.paragraph_format.space_before=Twips(200); h1.paragraph_format.space_after=Twips(200)
    h1.paragraph_format.line_spacing_rule=WD_LINE_SPACING.MULTIPLE; h1.paragraph_format.line_spacing=1.5
    h1.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
    h2=doc.styles['Heading 2']; style_fonts(h2,12,True); h2.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.line_spacing_rule=WD_LINE_SPACING.MULTIPLE; h2.paragraph_format.line_spacing=BODY
    h2.paragraph_format.space_before=Pt(0); h2.paragraph_format.space_after=Pt(8)
    t1=get_or_add_style(doc,'TOC1'); style_fonts(t1,11,True); dotted_tab(t1)
    t2=get_or_add_style(doc,'TOC2'); style_fonts(t2,10,False); dotted_tab(t2)

    # 预分配书签
    bid=1; toc=[]   # (level, text, bmname)
    for sec in data["sections"]:
        sn=f"_Toc_s{sec['no']}"; toc.append((1,f"{sec['no']}. {sec['title']}",sn,bid)); bid+=1
        for e in sec.get("entries",[]):
            en=f"_Toc_{sec['no']}_{str(e['no']).replace('.','_')}"
            toc.append((2,f"{e['no']} {e['title']}",en,bid)); bid+=1

    # 目录(开头直接是“目录”, 无标题/日期)
    font_run(para(doc,1.5,0,200).add_run("目录"),14,True)
    for level,text,bm,_ in toc:
        if level==1: toc_entry(doc,'TOC1',text,bm,11,True)
        else:        toc_entry(doc,'TOC2',text,bm,10,False)
    doc.add_page_break()

    # 正文
    bmmap={t[2]:t[3] for t in toc}
    for sec in data["sections"]:
        sn=f"_Toc_s{sec['no']}"
        heading_bm(doc,'Heading 1',f"{sec['no']}. {sec['title']}",14,sn,bmmap[sn])
        for e in sec.get("entries",[]):
            en=f"_Toc_{sec['no']}_{str(e['no']).replace('.','_')}"
            heading_bm(doc,'Heading 2',f"{e['no']} {e['title']}",12,en,bmmap[en])
            font_run(para(doc, after=160, align=WD_ALIGN_PARAGRAPH.JUSTIFY).add_run(e.get("abstract","")),12)
            cp=para(doc, after=160, align=WD_ALIGN_PARAGRAPH.LEFT)
            font_run(cp.add_run("【原文信息】："),12,True); font_run(cp.add_run(e.get("citation","")),12)
            font_run(para(doc, after=160).add_run(""),12)   # 子模块间空行
    doc.save(out); print("saved:",out)

if __name__=="__main__": main()
